from pandas import DataFrame
import pandas as pd
import numpy as np
from decimal import *
from docx import Document
from docx.shared import Inches
from django.shortcuts import render
from . import models
from . import batchprocess
from . import bof_singlecost
from django.shortcuts import render
from django.http import HttpResponse,HttpResponseRedirect,StreamingHttpResponse
import os
import json
import math


#暴力追溯时程序入口(论文实验用！)
def violent_analyse(request):
	print("start violent_analyse paper experiment")

	# fieldname_ch_list=['总吹氧耗量','出钢量','钢水温度','C','SI','MN']
	# fieldname_en_list=['SUM_BO_CSM','STEELWGT','FINAL_TEMP_VALUE','C','SI','MN']
	# filepath_list = ['e:/1636000_SUM_BO_CSM_100.docx','e:/1636000_STEELWGT_100.docx','e:/1636000_FINAL_TEMP_VALUE_100.docx','e:/1636000_C_100.docx','e:/1636000_SI_100.docx','e:/1636000_MN_100.docx']
	fieldname_ch_list=['钢渣']
	fieldname_en_list=['STEEL_SLAG']
	filepath_list = ['e:/1631230_STEEL_SLAG_100.docx']
	except_list = ['NB']
	for i in range(len(fieldname_ch_list)):
		fieldname_ch = fieldname_ch_list[i]
		fieldname_en = fieldname_en_list[i]
		filepath = filepath_list[i]
		document = Document()
		paragraph=document.add_paragraph()
		result={}
		#炉次		
		num=0
		for i in range(3000):
			prime_cost=str(1631230+i);
			single_result=violent_analyse_to(prime_cost,document,paragraph,filepath,fieldname_ch,fieldname_en)
			if single_result==None :
				continue
			else:
				for factor in single_result:
					if factor  not in result.keys():
						result[factor]=0
					result[factor]+=1

				num=num+1
				if num>=10:
					break
		print("追溯的异常炉次个数：",num)
		paragraph.add_run(str(result))
		document.save(filepath)


	contentVO={
		'title':'测试',
		'state':'success'		
	}				
	return HttpResponse(json.dumps(contentVO),content_type='application/json')	



def violent_analyse_to(prime_cost,document,paragraph,filepath,fieldname_ch,fieldname_en):
	# str_heatno='-------------------炉次号:'+prime_cost+'-------------------\n'      
	field_classification=['test','raw','material','product','alloy','quality']

	for k in range(0,1):
		# classify='【'+field_classification[k]+'】\n'
		# paragraph.add_run(classify)
		if field_classification[k]=='test':
			xasis_fieldname_ch=[fieldname_ch]
			xasis_fieldname_en=[fieldname_en]
			danwei=['Kg']
		elif field_classification[k]=='raw':
			#原料
			xasis_fieldname_ch=['铁水重量','生铁','废钢总和','大渣钢','自产废钢','重型废钢','中型废钢']
			xasis_fieldname_en=['MIRON_WGT','COLDPIGWGT','SCRAPWGT_COUNT','SCRAP_96053101','SCRAP_96052200','SCRAP_16010101','SCRAP_16020101']
			danwei=['Kg','Kg','Kg','Kg','Kg','Kg','Kg']
		elif field_classification[k]=='material':
			#物料
			xasis_fieldname_ch=['总吹氧消耗','氮气耗量','1#烧结矿','石灰石_40-70mm','萤石_FL80','增碳剂','低氮增碳剂','石灰','轻烧白云石']
			xasis_fieldname_en=['SUM_BO_CSM','N2CONSUME','L96020400','L12010302','L12010601','L12020201','L12020301','L96040100','L96040200']
			danwei=['NM3','NM3','Kg','Kg','Kg','Kg','Kg','Kg','Kg']
		elif field_classification[k]=='product':
			#产品
			xasis_fieldname_ch=['出钢量','转炉煤气','钢渣']
			xasis_fieldname_en=['STEELWGT','LDG_STEELWGT','STEEL_SLAG']
			danwei=['Kg','NM3','Kg']
		elif field_classification[k]=='alloy':
			#合金
			xasis_fieldname_ch=['硅铁_Si72-80%、AL≤2%(粒度10-60mm)','微铝硅铁_Si 72-80%、AL≤0.1%、Ti≤0.1%','硅锰合金_Mn 65-72%、Si 17-20%','高硅硅锰_Mn ≥60%、Si ≥27%','中碳铬铁']
			xasis_fieldname_en=['L13010101','L13010301','L13020101','L13020201','L13040400']
			danwei=['Kg','Kg','Kg','Kg','Kg']
		else:
			xasis_fieldname_single=['C','SI','MN','P','S','Fe','STEELWGT','FINAL_TEMP_VALUE']
			xaxis=['C含量','SI含量','MN含量','P含量','S含量','Fe含量','重量','温度']
			danwei=['%','%','%','%','%','%','Kg','℃']

		sqlVO={}
		sqlVO["db_name"]="l2own"
		sqlVO["sql"] = "SELECT " + ",".join(xasis_fieldname_en) + " from QG_USER.PRO_BOF_HIS_ALLFIELDS where heat_no= '"+prime_cost+"'"
		scrapy_records_single=models.BaseManage().direct_select_query_sqlVO(sqlVO)

		if len(scrapy_records_single) == 0:
			print(prime_cost+"炉次不存在")
			return None
		else:
			print(prime_cost+"炉次存在")


		yaxis_single=[]

		for i in range(len(xasis_fieldname_en)):
			value = scrapy_records_single[0].get(xasis_fieldname_en[i],None)
			# print(value)
			if value != None :
				scrapy_records_single[0][xasis_fieldname_en[i]] = float(value)
			else:
				value = 0
			yaxis_single.append(float(value))#由于在数据库查询时已经将空值设为了0，因此正常情况不会出现字段空值的情况

		print('字段名：',xasis_fieldname_en)
		print("字段实际值",yaxis_single)
		
		#计算单炉次质量分析字段偏离程度
		desired,offset_result_single=bof_singlecost.offset(xasis_fieldname_en,yaxis_single,'',True,3)
		offset_value_single=["%.2f%%"%(n*100) for n in list(offset_result_single)]
		offset_value_single_abs=["%.2f%%"%(abs(float(n))*100) for n in list(offset_result_single)]
		print('字段偏离度',offset_result_single)

		#分析字段偏离程度定性判断
		qualitative_offset_result=bof_singlecost.qualitative_offset(offset_result_single)
		# print('偏离度定性分析',qualitative_offset_result)

		#写入word
		xasis_fieldname_result_max=None
		for i in range(len(xasis_fieldname_en)):
			print("进行%s的追溯"%(xasis_fieldname_en[i]))
			xaxis_chinese=xasis_fieldname_ch[i];
			field=xasis_fieldname_en[i];
			single_value=yaxis_single[i];
			offset_value=offset_result_single[i];
			# offset_value_abs=offset_value_single_abs[i]
			# qualitative_offset_result_single=qualitative_offset_result[i];

			sql = "select DATA_ITEM_EN,IF_FIVENUMBERSUMMARY,NUMERICAL_LOWER_BOUND,NUMERICAL_UPPER_BOUND from QG_USER.PRO_BOF_HIS_ALLSTRUCTURE WHERE  DATA_ITEM_EN = '"+field+"'"
			sqlVO["sql"] = sql
			scrapy_records = models.BaseManage().direct_select_query_sqlVO(sqlVO)
			NUMERICAL_LOWER_BOUND=scrapy_records[0].get('NUMERICAL_LOWER_BOUND',None)#下限
			NUMERICAL_UPPER_BOUND=scrapy_records[0].get('NUMERICAL_UPPER_BOUND',None)#上限
			if single_value==0 or offset_value==None:#实际值为0的字段在暴力追溯中暂时不分析
				# str_des='本炉次'+prime_cost+'的'+xaxis_chinese+'数据异常！\n' 
				# paragraph.add_run(str_des)	
				# paragraph.add_run('\n')	
				continue
			if (field_classification[k]== 'material' or field_classification[k]== 'alloy' ) and  offset_value<0:
				continue
				
			if single_value<float(NUMERICAL_LOWER_BOUND) or single_value>float(NUMERICAL_UPPER_BOUND):#实际值在上下限范围之外，表示数据异常；实际值在上下限之间，但在最大最小值之外，在计算偏离程度时（offeset()）按照最大最小值计算
				# str_des='本炉次'+prime_cost+'的'+xaxis_chinese+'数据异常！\n' 
				# paragraph.add_run(str_des)	
				# paragraph.add_run('\n')	
				continue
			elif abs(float(offset_value))<=0.2:#偏离度小于10%的设定为正常
				# str_des='本炉次'+prime_cost+'的'+xaxis_chinese+qualitative_offset_result_single+',实际值为'+str(single_value)+danwei[i]+',偏离度为'+offset_value+'。\n'      
				# paragraph.add_run(str_des)
				continue

			En_to_Ch_result_score,xasis_fieldname_result_max,offset_result_nature,offset_value_single_cof,regression_coefficient_result=analy_cof(prime_cost,field,single_value,offset_value,'',True,3);		
			if xasis_fieldname_result_max != None:
				paragraph.add_run(prime_cost+' '+str(xaxis_chinese)+' '+str(En_to_Ch_result_score)+'\n')
			# if 	En_to_Ch_result_score==None:
			# 	# str_des='本炉次'+prime_cost+'的'+xaxis_chinese+qualitative_offset_result_single+',实际值为'+str(single_value)+danwei[i]+'，但进行回归分析时相关字段无数据！'
			# 	# paragraph.add_run(str_des)	
			# 	# paragraph.add_run('\n')	
			# 	continue
			# else:
				# # str_des='本炉次'+prime_cost+'的'+xaxis_chinese+qualitative_offset_result_single+',实际值为'+str(single_value)+danwei[i]+',偏离度为'+offset_value+'。通过数据相关性分析发现，导致该问题的原因是:\n'      
				# str_des=xaxis_chinese+qualitative_offset_result_single+offset_value_abs+'：'      
				
				# paragraph.add_run(str_des)	
				# for i in range(len(En_to_Ch_result_score)):
				# 	str_cause=En_to_Ch_result_score[i]+offset_result_nature[i]+offset_value_single_cof[i]+'；'
				# 	paragraph.add_run(str_cause) 
				# paragraph.add_run('\n')
			# paragraph.add_run('\n')
	# document.add_page_break()
	document.save(filepath)
	return 	xasis_fieldname_result_max	

from . import zhuanlu
def  analy_cof(prime_cost,field,single_value,offset_value,str_select,ifcache,whichcache):
	#从数据库读取相关字段并按照相关系数绝对值由大到小排序
	print('进行%s炉次下的%s字段的追溯'%(prime_cost,field))
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT * FROM pro_bof_his_relation_cof where OUTPUTFIELD='"+field +"'order by abs(COF) desc"
	#print(sqlVO["sql"])
	scrapy_records_relation=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	#print(scrapy_records)
	length_result1=len(scrapy_records_relation)
	if length_result1==0:#即无相关性字段
		return None,None,None,None,None


	xasis_fieldname=[]#字段英文名字数组
	correlation_coefficient=[]#字段相关系数值数组
	str_sql=''
	for i in range(length_result1):
		xasis_fieldname.append(scrapy_records_relation[i].get('MIDDLEFIELD', None))
		correlation_coefficient.append(scrapy_records_relation[i].get('COF', None))
		str_sql=str_sql+','+scrapy_records_relation[i].get('MIDDLEFIELD', None)
	
	#取相关字段实际值
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO" +str_sql+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS where HEAT_NO='"+prime_cost+"'";
	print(sqlVO["sql"])
	scrapy_records_actual=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	yaxis=[]#各相关性字段实际值
	vaildnumber = 0#有数据的相关字段个数，如果有数据的相关字段太少，则说明该炉次的有效字段太少，不足以追溯
	for i in range(length_result1):
		value = scrapy_records_actual[0].get(xasis_fieldname[i],None)
		if value != None :
			vaildnumber +=1
			scrapy_records_actual[0][xasis_fieldname[i]] = Decimal(value)
		else:
			value = 0#将空值None暂时以0填充
		yaxis.append(value)

	if vaildnumber<8:
		return None,None,None,None,None #有数据的当前炉次的相关字段太少，没有追溯的必要，追溯误差大
	desired,offset_degree=bof_singlecost.offset(xasis_fieldname,yaxis,'',True,3)
	# print("相关字段名字：",xasis_fieldname)
	# print("相关字段实际值：",yaxis)
	# print("相关性：",correlation_coefficient)
	# print("相关字段偏离程度：",offset_degree)

	xasis_fieldname_result=[]#字段英文名字数组
	correlation_coefficient_result=[]#字段回归系数值数组
	offset_degree_result=[]#偏离程度值

	j=0#标志位，表示当前筛选后的有效字段个数
	if float(offset_value)>=0:
		for i in range(length_result1):
			if j>=8:#取筛选后相关性最大的八个因素字段
				break
			#去除一些冗余字段
			# if abs(float(correlation_coefficient[i]))>0.19 and abs(float(correlation_coefficient[i]))<0.3 or abs(float(correlation_coefficient[i]))>0.5:
			# 	continue
			if offset_degree[i]==None  or abs(offset_degree[i])<=0.1 or yaxis[i]==0 :
				continue
			if float(correlation_coefficient[i]) * float(offset_degree[i]) >=0:
				j+=1
				xasis_fieldname_result.append(xasis_fieldname[i])
				correlation_coefficient_result.append(correlation_coefficient[i])
				offset_degree_result.append(offset_degree[i])
	else:
		for i in range(length_result1):
			if j>=8:
				break
			# if abs(float(correlation_coefficient[i]))>0.19 and abs(float(correlation_coefficient[i]))<0.3 or abs(float(correlation_coefficient[i]))>0.5:
			# 	continue
			if offset_degree[i]==None  or abs(offset_degree[i])<=0.1 or yaxis[i]==0:
				continue
			if float(correlation_coefficient[i]) * float(offset_degree[i]) <=0:
				j+=1
				xasis_fieldname_result.append(xasis_fieldname[i])
				correlation_coefficient_result.append(correlation_coefficient[i])
				offset_degree_result.append(offset_degree[i])
	
	#正负相关筛选后的相关字段个数
	print('正负相关筛选后的相关字段个数',len(xasis_fieldname_result))
	print("相关字段英文名字：",xasis_fieldname_result)

	if len(xasis_fieldname_result) <=3:#正负相关筛选后的相关字段个数太少
		return None,None,None,None,None

	#计算回归系数:regression_coefficient[0]表示各回归值，regression_coefficient[1]表示截距;1表示在regression中进行标准化，0表示不进行标准化
	regression_coefficient=batchprocess.regression(field,xasis_fieldname_result,1)
	if regression_coefficient== False:
		return None,None,None,None,None

	#查询转炉工序字段名中英文对照
	ana_result={}
	ana_result=zhuanlu.PRO_BOF_HIS_ALLFIELDS
	En_to_Ch_result=[]
	for i in range(len(xasis_fieldname_result)):
		En_to_Ch_result.append(ana_result[xasis_fieldname_result[i]])
	# print("字段中文名字")
	# print(En_to_Ch_result)

	# print("8个字段英文名字：")
	# print(xasis_fieldname_result)
	# print("相关性系数")
	# print(correlation_coefficient_result)
	# print("回归系数")
	# print(regression_coefficient[0])
	# print("偏离程度")
	# print(offset_degree_result)

	#zip压缩:中文名、英文名、偏离程度、回归系数;注：python3之后zip函数返回的是迭代值，需要list强转
	L_zip=list(zip(En_to_Ch_result,xasis_fieldname_result,offset_degree_result,regression_coefficient[0],correlation_coefficient_result))
	# print('L_zip:',list(L_zip))
	#按照回归系数进行排序(从大到小)
	L_zip.sort(key=lambda x:x[3],reverse=True)
	# print('按照回归系数排序后的字段：',L_zip)
	#取回归系数前三的字段
	L_zip=L_zip[0:2]
	# print('取前两个字段时的实际字段个数：',len(L_zip))
	#解压缩
	L_unzip=list(zip(*L_zip))
	En_to_Ch_result_max=L_unzip[0]#中文名
	xasis_fieldname_result_max=L_unzip[1]#英文名
	offset_degree_result_max=L_unzip[2]#偏离程度
	regression_coefficient_max=L_unzip[3]#回归系数（权重）

	# # '''
	# #判断字段的顺序（对筛选后的三个字段再根据字段再表结构中的顺序进行排序）
	# ana_result_score={}
	# ana_result_score=zhuanlu.PRO_BOF_HIS_ALLFIELDS_SCORE
	# En_to_Ch_result_score=[]
	# for i in range(len(xasis_fieldname_result_max)):
	# 	En_to_Ch_result_score.append(ana_result_score[xasis_fieldname_result_max[i]][1])
	# # print('表结构中字段的顺序',En_to_Ch_result_score)

	# #zip压缩：中文名、英文名、偏离程度、字段顺序编号
	# L_zip = list(zip(En_to_Ch_result_max,xasis_fieldname_result_max,offset_degree_result_max,En_to_Ch_result_score,regression_coefficient_max))
	# #按照字段顺序进行排序（从小到大）
	# L_zip.sort(key=lambda x:x[3])
	# print('追溯结果的三个字段',L_zip)
	# #解压缩
	# L_unzip=list(zip(*L_zip))

	# En_to_Ch_result_max=L_unzip[0]#中文名
	# xasis_fieldname_result_max=L_unzip[1]#英文名
	# offset_degree_result_max=L_unzip[2]#偏离程度
	# En_to_Ch_result_score_max=L_unzip[3]#字段序号
	# regression_coefficient_max=L_unzip[4]#回归系数（权重）
	# # '''

	offset_degree_result_max_final=["%.2f%%"%(abs(float(n))*100) for n in list(offset_degree_result_max)]#将偏离程度值转化为保留两位小数的百分数
	#简单定性判断偏离程度（偏高、偏低）
	pos_float=[ float(n) for n in list(offset_degree_result_max)]
	offset_result_nature=simple_offset(pos_float)



	# print('分析字段：',field)
	# print('分析字段偏离程度定性判断')
	# print(qualitative_offset_result[0])
	# print('分析字段偏离程度')
	# print(offset_value_single[0])
	# print('按操作排序后回归系数中文名')
	# print(En_to_Ch_result_score)
	# print("简单定性判断")
	# print(offset_result_nature)
	# print('按操作排序后字段偏离程度')
	# print(offset_value_single_cof)
	return En_to_Ch_result_max,xasis_fieldname_result_max,offset_result_nature,offset_degree_result_max_final,regression_coefficient_max


#对相关性字段展示做简单定性分析	正数为偏高负数为偏低
def simple_offset(offset_result_final):
	#偏离程度定性标准，例如-10%~10%为正常，10%~30%为偏高，30%为高，40%以上为数据异常/极端数据
	offset_result_nature=[]#相关字段定性分析
	for i in range(len(offset_result_final)):
		if float(offset_result_final[i]>0):
			offset_result_nature.append('偏高')
		else:
			offset_result_nature.append('偏低')	
	return  offset_result_nature



# #暴力求解
# if __name__ == '__main__':
# 	violent_analyse()